"""
Router de ingesta de documentos.

Endpoints:
  POST /api/ingest              — síncrono (retorna cuando completa)
  POST /api/ingest/async        — asíncrono (retorna job_id inmediatamente)
  GET  /api/ingest/status/{id}  — estado del job asíncrono

Flujo (ambos endpoints):
  1. Leer y detectar MIME (PDF / TXT / DOCX)
  2. Extraer texto
  3. Chunking + indexar en Qdrant (colección por usuario)
  4. Generar resumen con LLM
  5. Guardar metadata en PostgreSQL (user_documents)
  6. Backup en MinIO (best-effort)

Migrado desde backend-ia/main.py → ingest_data().
"""
from __future__ import annotations

import logging
import os
import uuid

from fastapi import APIRouter, BackgroundTasks, Depends, File, HTTPException, UploadFile

from interfaces.api.auth import get_current_user

logger = logging.getLogger("interfaces.api.routers.ingest")

router = APIRouter(prefix="/api", tags=["ingest"])

# Singleton LLM para resúmenes de documentos
_ingest_llm = None


def _get_ingest_llm():
    global _ingest_llm
    if _ingest_llm is None:
        from agents.base.llm_factory import get_chat_llm
        _ingest_llm = get_chat_llm(timeout=60)
    return _ingest_llm


def _extract_docx_text(content: bytes) -> str:
    """Extrae texto de un archivo DOCX."""
    import io
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        logger.warning(f"[ingest] Error extrayendo DOCX: {e}")
        return ""


def _sanitize_email(email: str) -> str:
    """Convierte email en nombre válido para bucket MinIO."""
    return email.replace("@", "_at_").replace(".", "_")


@router.post("/ingest")
async def ingest_document(
    file: UploadFile = File(...),
    user: str = Depends(get_current_user),
):
    """
    Sube y procesa un documento (PDF, TXT, DOCX, MD).
    Lo indexa en Qdrant y guarda metadata en PostgreSQL.

    Returns:
        { doc_id, filename, summary, chunks }
    """
    temp_path: str | None = None
    content = await file.read()

    # ── 1. Detectar MIME ──────────────────────────────────────────────────────
    try:
        import magic
        mime = magic.from_buffer(content, mime=True)
    except Exception:
        raise HTTPException(status_code=400, detail="No se pudo determinar el tipo de archivo.")

    fname = (file.filename or "").lower()
    if fname.endswith(".docx"):
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if fname.endswith(".md"):
        mime = "text/plain"

    # ── 2. Extraer texto ──────────────────────────────────────────────────────
    from langchain_core.documents import Document as LCDocument

    full_text = ""
    documents: list[LCDocument] = []

    if mime == "application/pdf":
        temp_path = f"/tmp/{uuid.uuid4()}-{file.filename}"  # nosec B108 — uuid4 hace el nombre único e impredecible
        with open(temp_path, "wb") as buf:
            buf.write(content)
        try:
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(temp_path)
            documents = loader.load()
            full_text = "\n".join(d.page_content for d in documents)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error leyendo PDF: {e}")

    elif mime == "text/plain":
        full_text = content.decode("utf-8", errors="replace")
        documents = [LCDocument(page_content=full_text)]

    elif "wordprocessingml" in mime or fname.endswith(".docx"):
        full_text = _extract_docx_text(content)
        documents = [LCDocument(page_content=full_text)]

    else:
        raise HTTPException(
            status_code=400,
            detail=f"Tipo de archivo no soportado: '{mime}'. Usa PDF, TXT, DOCX o MD.",
        )

    if not full_text.strip():
        raise HTTPException(status_code=422, detail="El documento está vacío o no se pudo extraer texto.")

    # ── 3. Chunking + indexar en Qdrant ───────────────────────────────────────
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        from agents.researcher.rag_retriever import (
            _get_qdrant_client,
            get_user_vectorstore,
            sanitize_email,
        )

        # P7-003: Document versioning — eliminar chunks previos del mismo filename
        try:
            qdrant_client = _get_qdrant_client()
            collection_name = sanitize_email(user)
            if qdrant_client.collection_exists(collection_name):
                # Scroll + delete por substring en Python (sin FTS index requerido)
                fname_lower = (file.filename or "").lower()
                scroll_result, _ = qdrant_client.scroll(
                    collection_name=collection_name,
                    limit=1000,
                    with_payload=True,
                    with_vectors=False,
                )
                ids_to_delete = [
                    p.id for p in scroll_result
                    if fname_lower in (
                        (p.payload or {}).get("metadata", {}).get("filename") or ""
                    ).lower()
                    or fname_lower in (
                        (p.payload or {}).get("metadata", {}).get("source") or ""
                    ).lower()
                ]
                if ids_to_delete:
                    qdrant_client.delete(
                        collection_name=collection_name,
                        points_selector=ids_to_delete,
                    )
                    logger.info(f"[ingest] Versioning: eliminados {len(ids_to_delete)} chunks previos de '{file.filename}'")
        except Exception as ve:
            logger.warning(f"[ingest] Versioning delete falló (no crítico): {ve}")

        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = splitter.split_documents(documents)
        # Asegurar que cada chunk lleva el filename limpio en metadata
        for chunk in chunks:
            chunk.metadata["filename"] = file.filename
        vs = get_user_vectorstore(user)
        vs.add_documents(chunks)
        n_chunks = len(chunks)
        logger.info(f"[ingest] {n_chunks} chunks indexados en Qdrant para '{user}'")
    except Exception as e:
        logger.error(f"[ingest] Error indexando en Qdrant: {e}")
        raise HTTPException(status_code=500, detail=f"Error indexando documento: {e}")

    # ── 4. Generar resumen con LLM ────────────────────────────────────────────
    summary = ""
    try:
        llm = _get_ingest_llm()
        _resp = llm.invoke(
            f"Resume el siguiente documento en 2-3 oraciones en español:\n\n{full_text[:3000]}"
        )
        summary = (_resp.content if hasattr(_resp, "content") else str(_resp)).strip()
    except Exception as e:
        summary = f"Documento procesado ({n_chunks} fragmentos)."
        logger.warning(f"[ingest] Error generando resumen: {e}")

    # ── 5. Metadata en PostgreSQL ─────────────────────────────────────────────
    doc_id = None
    try:
        from storage.postgres.client import get_connection
        with get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "INSERT INTO user_documents (user_id, doc_type, summary, raw_analysis) "
                    "VALUES (%s, %s, %s, %s) RETURNING id",
                    (user, file.filename, summary, full_text[:10000]),
                )
                doc_id = cur.fetchone()[0]
            conn.commit()
    except Exception as e:
        logger.warning(f"[ingest] No se guardó metadata en PostgreSQL: {e}")

    # ── 6. Backup en MinIO (best-effort) ─────────────────────────────────────
    try:
        import io

        from storage.minio.client import get_client
        minio = get_client()
        bucket = _sanitize_email(user).replace("_", "-")[:63]
        if not minio.bucket_exists(bucket):
            minio.make_bucket(bucket)
        file_data = content if temp_path is None else open(temp_path, "rb").read()
        minio.put_object(
            bucket,
            file.filename,
            io.BytesIO(file_data),
            length=len(file_data),
            content_type=mime,
        )
        logger.info(f"[ingest] Backup MinIO: {bucket}/{file.filename}")
    except Exception as e:
        logger.warning(f"[ingest] MinIO backup falló (no crítico): {e}")

    # ── Cleanup ───────────────────────────────────────────────────────────────
    if temp_path and os.path.exists(temp_path):
        os.remove(temp_path)

    return {
        "doc_id":   doc_id,
        "filename": file.filename,
        "summary":  summary,
        "chunks":   n_chunks,
    }


# ── Job tracking helpers ──────────────────────────────────────────────────────

_JOB_TTL = 3600  # 1 hora — tiempo máximo que se guarda el estado del job en Redis


def _job_key(job_id: str) -> str:
    return f"ingest_job:{job_id}"


def _set_job_status(job_id: str, status: str, **extra) -> None:
    """Actualiza el estado del job en Redis."""
    import json as _json
    try:
        from storage.redis.client import get_redis_client
        rc = get_redis_client()
        payload = {"status": status, "job_id": job_id, **extra}
        rc.setex(_job_key(job_id), _JOB_TTL, _json.dumps(payload))
    except Exception as exc:
        logger.warning(f"[ingest] No se pudo actualizar job {job_id}: {exc}")


def _notify_ingest_done(
    user: str,
    filename: str,
    chunks: int,
    success: bool,
    error: str = "",
) -> None:
    """
    Notifica al usuario por WhatsApp cuando su ingestión asíncrona termina.
    Busca el número WhatsApp del usuario en la tabla de identidades.
    No-op si el usuario no tiene WhatsApp registrado o el bridge no está configurado.
    """
    _wa_url = os.environ.get("WHATSAPP_BRIDGE_URL", "http://whatsapp-bridge-service:3000")
    if not _wa_url:
        return
    try:
        # Buscar número WhatsApp del usuario en identidades
        from storage.postgres.client import get_connection
        with get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT identifier FROM user_identities "
                    "WHERE user_id = %s AND identity_type = 'whatsapp' LIMIT 1",
                    (user,),
                )
                row = cur.fetchone()
        if not row:
            return
        phone = row[0]

        if success:
            msg = (
                f"✅ *Documento procesado correctamente*\n"
                f"📄 Archivo: `{filename}`\n"
                f"🔢 Fragmentos indexados: {chunks}\n"
                f"Ya puedes hacer preguntas sobre este documento."
            )
        else:
            msg = (
                f"❌ *Error procesando documento*\n"
                f"📄 Archivo: `{filename}`\n"
                f"⚠️ {error}"
            )

        import requests as _req
        _req.post(
            f"{_wa_url}/send",
            json={"phone": phone, "message": msg},
            timeout=8,
        )
        logger.info(f"[ingest] Notificación WhatsApp enviada a {phone!r}")
    except Exception as exc:
        logger.debug(f"[ingest] WhatsApp notify failed (non-critical): {exc}")


def _run_ingest_job(
    job_id: str,
    user: str,
    filename: str,
    content: bytes,
    mime: str,
) -> None:
    """
    Ejecuta el pipeline de ingestión en segundo plano.
    Actualiza el estado del job en Redis a lo largo del proceso.
    """
    import io

    _set_job_status(job_id, "processing", filename=filename)

    try:
        # ── Extraer texto ──────────────────────────────────────────────────
        from langchain_core.documents import Document as LCDocument

        full_text = ""
        documents: list[LCDocument] = []
        temp_path = None

        if mime == "application/pdf":
            temp_path = f"/tmp/{uuid.uuid4()}-{filename}"  # nosec B108 — uuid4 hace el nombre único e impredecible
            with open(temp_path, "wb") as buf:
                buf.write(content)
            try:
                from langchain_community.document_loaders import PyPDFLoader
                loader = PyPDFLoader(temp_path)
                documents = loader.load()
                full_text = "\n".join(d.page_content for d in documents)
            finally:
                if temp_path and os.path.exists(temp_path):
                    os.remove(temp_path)
        elif mime == "text/plain":
            full_text = content.decode("utf-8", errors="replace")
            documents = [LCDocument(page_content=full_text)]
        else:
            full_text = _extract_docx_text(content)
            documents = [LCDocument(page_content=full_text)]

        if not full_text.strip():
            _set_job_status(job_id, "failed", error="El documento está vacío o no se pudo extraer texto.")
            return

        # ── Chunking + Qdrant ─────────────────────────────────────────────
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        from agents.researcher.rag_retriever import get_user_vectorstore
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = splitter.split_documents(documents)
        for chunk in chunks:
            chunk.metadata["filename"] = filename
        vs = get_user_vectorstore(user)
        vs.add_documents(chunks)
        n_chunks = len(chunks)

        # ── Resumen LLM ───────────────────────────────────────────────────
        summary = ""
        try:
            llm = _get_ingest_llm()
            _resp = llm.invoke(
                f"Resume el siguiente documento en 2-3 oraciones en español:\n\n{full_text[:3000]}"
            )
            summary = (_resp.content if hasattr(_resp, "content") else str(_resp)).strip()
        except Exception:
            summary = f"Documento procesado ({n_chunks} fragmentos)."

        # ── PostgreSQL ────────────────────────────────────────────────────
        doc_id = None
        try:
            from storage.postgres.client import get_connection
            with get_connection() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        "INSERT INTO user_documents (user_id, doc_type, summary, raw_analysis) "
                        "VALUES (%s, %s, %s, %s) RETURNING id",
                        (user, filename, summary, full_text[:10000]),
                    )
                    doc_id = cur.fetchone()[0]
                conn.commit()
        except Exception as exc:
            logger.warning(f"[ingest] PostgreSQL insert failed (async job {job_id}): {exc}")

        # ── MinIO backup ──────────────────────────────────────────────────
        try:
            from storage.minio.client import get_client
            minio = get_client()
            bucket = _sanitize_email(user).replace("_", "-")[:63]
            if not minio.bucket_exists(bucket):
                minio.make_bucket(bucket)
            minio.put_object(bucket, filename, io.BytesIO(content), len(content), content_type=mime)
        except Exception as exc:
            logger.warning(f"[ingest] MinIO backup failed (async job {job_id}): {exc}")

        _set_job_status(
            job_id, "completed",
            filename=filename,
            doc_id=str(doc_id) if doc_id else None,
            summary=summary,
            chunks=n_chunks,
        )
        _notify_ingest_done(user, filename, n_chunks, success=True)

    except Exception as exc:
        logger.error(f"[ingest] Async job {job_id} failed: {exc}", exc_info=True)
        _set_job_status(job_id, "failed", filename=filename, error=str(exc)[:300])
        _notify_ingest_done(user, filename, 0, success=False, error=str(exc)[:200])


# ── Async endpoints ───────────────────────────────────────────────────────────

@router.post("/ingest/async", status_code=202)
async def ingest_document_async(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    user: str = Depends(get_current_user),
):
    """
    Inicia la ingestión en segundo plano y retorna inmediatamente con un job_id.
    Útil para documentos grandes (PDF de muchas páginas) que tardan > 30s.

    Returns:
        { job_id, status: "queued", filename }
    """
    content = await file.read()
    fname = (file.filename or "").lower()

    # MIME detection
    try:
        import magic
        mime = magic.from_buffer(content, mime=True)
    except Exception:
        raise HTTPException(status_code=400, detail="No se pudo determinar el tipo de archivo.")

    if fname.endswith(".docx"):
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if fname.endswith(".md"):
        mime = "text/plain"

    supported = {"application/pdf", "text/plain",
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    if mime not in supported and not fname.endswith(".docx"):
        raise HTTPException(status_code=400, detail=f"Tipo no soportado: '{mime}'")

    job_id = str(uuid.uuid4())
    _set_job_status(job_id, "queued", filename=file.filename)
    background_tasks.add_task(_run_ingest_job, job_id, user, file.filename, content, mime)

    return {"job_id": job_id, "status": "queued", "filename": file.filename}


@router.get("/ingest/status/{job_id}")
async def get_ingest_status(
    job_id: str,
    user: str = Depends(get_current_user),
):
    """
    Consulta el estado de un job de ingestión asíncrona.

    Returns:
        { job_id, status: "queued"|"processing"|"completed"|"failed", ... }
    """
    import json as _json
    try:
        from storage.redis.client import get_redis_client
        rc = get_redis_client()
        raw = rc.get(_job_key(job_id))
        if not raw:
            raise HTTPException(status_code=404, detail="Job no encontrado o expirado.")
        return _json.loads(raw)
    except HTTPException:
        raise
    except Exception as exc:
        logger.error(f"[ingest] status lookup failed for {job_id}: {exc}")
        raise HTTPException(status_code=500, detail="Error consultando estado del job.")
